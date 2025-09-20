import os
from docx import Document
from flask_mail import Message
from flask import current_app
from werkzeug.utils import secure_filename


def fill_word_template(template_path, output_path, context):
    """
    Llena un archivo Word (.docx) usando un diccionario context.
    Reemplaza {{campo}} en el texto por el valor correspondiente.
    """
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, value in context.items():
            p.text = p.text.replace(f'{{{{{key}}}}}', str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
    doc.save(output_path)
    return output_path


def send_email_with_attachment(mail, subject, recipients, body, attachment_path):
    msg = Message(subject, recipients=recipients, body=body)
    with open(attachment_path, 'rb') as f:
        filename = os.path.basename(attachment_path)
        msg.attach(
            filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f.read())
    mail.send(msg)
